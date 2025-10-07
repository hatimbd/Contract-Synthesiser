from docx import Document
from rules import COLOR_MAP, RULES

def get_highlight_color(cell):
    """Retourne la couleur dominante de surlignage trouvée dans une cellule."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.font.highlight_color:
                try:
                    return run.font.highlight_color.name
                except:
                    return str(run.font.highlight_color)
    return None


def extract_cell_changes(filepath):
    """
    Extrait les modifications cellule par cellule du tableau Word.
    Chaque cellule surlignée donne lieu à une entrée : Log_ID, colonne, valeur, action.
    """
    doc = Document(filepath)
    table = doc.tables[0]
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    changes = []

    for row in table.rows[1:]:
        row_data = {headers[i]: row.cells[i].text.strip() for i in range(len(headers))}
        log_id = row_data.get("Log_ID", "").strip()

        # Parcours cellule par cellule
        for i, cell in enumerate(row.cells):
            color = get_highlight_color(cell)
            if not color:
                continue

            mapped_color = COLOR_MAP.get(color)
            action = RULES.get(mapped_color, "NONE")

            if action != "NONE":
                changes.append({
                    "Log_ID": log_id,
                    "column": headers[i],
                    "new_value": cell.text.strip(),
                    "action": action
                })

    return changes
